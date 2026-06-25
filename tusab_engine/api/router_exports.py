# Copyright (c) 2026 CriAugu — CNPJ 65.131.075/0001-57
"""
Endpoints de export — feature Pro.

  POST /export/base           → ZIP com cerebro/ + CSVs de gestão
  POST /export/historico      → Markdown do histórico de chat de um canal
  POST /export/resumo-canal   → DOCX com mensagens do assistente + fontes
  POST /export/tabela-videos  → XLSX com dados do CSV de gestão do canal
  POST /export/relatorio-pdf  → PDF com pares de Q&A do histórico
"""

import os
import io
import json
import zipfile
import shutil
import tempfile
from datetime import datetime

from fastapi import APIRouter, UploadFile, File
from fastapi.responses import StreamingResponse, JSONResponse
from pydantic import BaseModel, Field

import motor_tusab
from tusab_engine.state import state
from tusab_engine.storage import NEURAL_DIR, GESTAO_DIR, DATA_DIR, gestao_canal_dir, salvar_json_atomico

router = APIRouter()


# ── Export da base (ZIP) ──────────────────────────────────────────────────────

@router.post("/export/base")
def export_base():
    """Gera ZIP com cerebro/ + CSVs de gestão e serve como download."""
    buf = io.BytesIO()

    with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zf:
        # cerebro/
        if os.path.exists(NEURAL_DIR):
            for root, _, files in os.walk(NEURAL_DIR):
                for fname in files:
                    fpath = os.path.join(root, fname)
                    arcname = os.path.relpath(fpath, os.path.dirname(NEURAL_DIR))
                    zf.write(fpath, arcname)

        # gestao/ já está dentro de cerebro/{prefixo}/gestao/ — incluído no walk acima

    buf.seek(0)
    ts = datetime.now().strftime('%Y%m%d_%H%M')
    filename = f"tusab_base_{ts}.zip"

    return StreamingResponse(
        buf,
        media_type="application/zip",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )


# ── Export do histórico de chat ───────────────────────────────────────────────

class ExportHistoricoRequest(BaseModel):
    canal_nome: str = Field(default="", max_length=120)


@router.post("/export/historico")
def export_historico(req: ExportHistoricoRequest):
    """Exporta o histórico de chat de um canal como Markdown."""
    canal = req.canal_nome or state.stats.get("canal_nome", "") or "chat"

    with state.hist_lock:
        hist = list(state.chat_histories.get(canal, []))

    if not hist:
        return JSONResponse({"error": True, "message": "Nenhum histórico encontrado para este canal."})

    linhas = [f"# Histórico de Chat — @{canal}", f"_Exportado em {datetime.now().strftime('%d/%m/%Y %H:%M')}_\n"]
    for msg in hist:
        role = msg.get("role", "")
        content = msg.get("content", "")
        if role == "user":
            linhas.append(f"**Você:** {content}\n")
        elif role == "assistant":
            linhas.append(f"**Tusab:** {content}\n")
            fontes = msg.get("fontes", [])
            if fontes:
                linhas.append("_Fontes:_")
                for f in fontes:
                    titulo = f.get("titulo") or f.get("arquivo", "")
                    link = f.get("link", "")
                    if link:
                        linhas.append(f"- [{titulo}]({link})")
                    else:
                        linhas.append(f"- {titulo}")
                linhas.append("")

    md = "\n".join(linhas)
    buf = io.BytesIO(md.encode("utf-8"))
    ts = datetime.now().strftime('%Y%m%d_%H%M')
    filename = f"tusab_historico_{canal}_{ts}.md"

    return StreamingResponse(
        buf,
        media_type="text/markdown",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )


# ── Export de resumo do canal (DOCX) ─────────────────────────────────────────

class ExportResumoCanalRequest(BaseModel):
    canal_nome: str = Field(default="", max_length=120)
    mensagens:  list = Field(default_factory=list)  # histórico do frontend (fallback ao state)


@router.post("/export/resumo-canal")
def export_resumo_canal(req: ExportResumoCanalRequest):
    """Exporta as respostas do assistente para um canal como documento Word (.docx)."""
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        return JSONResponse({
            "error": True,
            "message": "Dependência não instalada: python-docx. Execute: pip install python-docx"
        })

    canal = req.canal_nome or state.stats.get("canal_nome", "") or "chat"

    # Prioriza mensagens enviadas pelo frontend; fallback ao histórico server-side
    if req.mensagens:
        hist = req.mensagens
    else:
        with state.hist_lock:
            hist = list(state.chat_histories.get(canal, []))

    if not hist:
        return JSONResponse({"error": True, "message": "Sem histórico para exportar"})

    doc = Document()

    # Título e cabeçalho
    doc.add_heading(f"Resumo — @{canal}", level=0)
    subtitulo = doc.add_paragraph("Gerado pelo Tusab")
    subtitulo.runs[0].italic = True
    doc.add_paragraph(f"Data: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    doc.add_paragraph("")

    for msg in hist:
        if msg.get("role") != "assistant":
            continue

        content = msg.get("content", "").strip()
        if not content:
            continue

        heading = doc.add_paragraph("Tusab:", style="Heading 2")

        doc.add_paragraph(content)

        fontes = msg.get("fontes", [])
        if fontes:
            doc.add_paragraph("Fontes:", style="Heading 3")
            for fonte in fontes:
                titulo = fonte.get("titulo") or fonte.get("arquivo", "")
                link = fonte.get("link", "")
                bullet = doc.add_paragraph(style="List Bullet")
                run = bullet.add_run(f"{titulo}" + (f" — {link}" if link else ""))
                run.font.size = Pt(10)

        doc.add_paragraph("")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    ts = datetime.now().strftime('%Y%m%d_%H%M')
    filename = f"tusab_resumo_{canal}_{ts}.docx"

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )


# ── Export de tabela de vídeos (XLSX) ────────────────────────────────────────

class ExportTabelaVideosRequest(BaseModel):
    canal: str = Field(default="", max_length=120)


@router.post("/export/tabela-videos")
def export_tabela_videos(req: ExportTabelaVideosRequest):
    """Exporta o CSV de gestão de um canal como planilha Excel (.xlsx)."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font
    except ImportError:
        return JSONResponse({
            "error": True,
            "message": "Dependência não instalada: openpyxl. Execute: pip install openpyxl"
        })

    canal = req.canal or state.stats.get("canal_nome", "") or ""
    if not canal:
        return JSONResponse({"error": True, "message": "Nome do canal não informado."})

    csv_path = os.path.join(gestao_canal_dir(canal), f"{canal}_base.csv")
    if not os.path.exists(csv_path):
        return JSONResponse({"error": True, "message": f"CSV não encontrado para o canal '{canal}'."})

    import csv

    with open(csv_path, newline="", encoding="utf-8") as fh:
        reader = csv.reader(fh)
        rows = list(reader)

    if not rows:
        return JSONResponse({"error": True, "message": "CSV vazio."})

    wb = Workbook()
    ws = wb.active
    ws.title = f"@{canal}"[:31]  # openpyxl limita nome de aba a 31 caracteres

    header = rows[0] if rows else []
    bold_font = Font(bold=True)

    for col_idx, cell_value in enumerate(header, start=1):
        cell = ws.cell(row=1, column=col_idx, value=cell_value)
        cell.font = bold_font

    for row_idx, row in enumerate(rows[1:], start=2):
        for col_idx, cell_value in enumerate(row, start=1):
            ws.cell(row=row_idx, column=col_idx, value=cell_value)

    # Congelar linha do cabeçalho
    ws.freeze_panes = "A2"

    # Largura automática aproximada
    for col_idx, col_header in enumerate(header, start=1):
        col_letter = ws.cell(row=1, column=col_idx).column_letter
        approx_width = min(len(str(col_header)) * 1.2 + 2, 50)
        ws.column_dimensions[col_letter].width = approx_width

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)

    ts = datetime.now().strftime('%Y%m%d_%H%M')
    filename = f"tusab_videos_{canal}_{ts}.xlsx"

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )


# ── Export de relatório de pesquisa (PDF) ─────────────────────────────────────

class ExportRelatorioPdfRequest(BaseModel):
    canal_nome: str = Field(default="", max_length=120)
    mensagens:  list = Field(default_factory=list)  # histórico do frontend (fallback ao state)


@router.post("/export/relatorio-pdf")
def export_relatorio_pdf(req: ExportRelatorioPdfRequest):
    """Exporta os pares de Q&A do histórico de um canal como PDF (.pdf)."""
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.units import mm
    except ImportError:
        return JSONResponse({
            "error": True,
            "message": "Recurso de exportação PDF não disponível. Reinstale a aplicação ou contacte o suporte em tusab@tusab.solutions"
        })

    canal = req.canal_nome or state.stats.get("canal_nome", "") or "chat"

    # Prioriza mensagens enviadas pelo frontend; fallback ao histórico server-side
    if req.mensagens:
        hist = req.mensagens
    else:
        with state.hist_lock:
            hist = list(state.chat_histories.get(canal, []))

    if not hist:
        return JSONResponse({"error": True, "message": "Sem histórico para exportar"})

    buf = io.BytesIO()

    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=20 * mm,
        rightMargin=20 * mm,
        topMargin=20 * mm,
        bottomMargin=20 * mm,
    )

    styles = getSampleStyleSheet()
    story = []

    # Título e cabeçalho
    story.append(Paragraph(f"Relatório de Pesquisa — @{canal}", styles["Title"]))
    story.append(Spacer(1, 4 * mm))
    story.append(Paragraph("Gerado pelo Tusab", styles["Italic"]))
    story.append(Paragraph(f"Data: {datetime.now().strftime('%d/%m/%Y %H:%M')}", styles["Normal"]))
    story.append(Spacer(1, 8 * mm))

    # Pares Q&A
    pending_question = None
    for msg in hist:
        role = msg.get("role", "")
        content = msg.get("content", "").strip()

        if role == "user":
            pending_question = content

        elif role == "assistant":
            if pending_question:
                story.append(Paragraph(f"<b>Pergunta:</b> {pending_question}", styles["Normal"]))
                story.append(Spacer(1, 2 * mm))
                pending_question = None

            story.append(Paragraph(f"<b>Tusab:</b> {content}", styles["Normal"]))

            fontes = msg.get("fontes", [])
            if fontes:
                story.append(Spacer(1, 2 * mm))
                for fonte in fontes:
                    titulo = fonte.get("titulo") or fonte.get("arquivo", "")
                    link = fonte.get("link", "")
                    texto_fonte = f"&nbsp;&nbsp;• {titulo}" + (f" — {link}" if link else "")
                    story.append(Paragraph(f'<font size="9">{texto_fonte}</font>', styles["Normal"]))

            story.append(Spacer(1, 6 * mm))

    doc.build(story)
    buf.seek(0)

    ts = datetime.now().strftime('%Y%m%d_%H%M')
    filename = f"tusab_relatorio_{canal}_{ts}.pdf"

    return StreamingResponse(
        buf,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )


# ── Export de base compartilhável (.tusab) ────────────────────────────────────

_INDEX_DIR = os.path.join(DATA_DIR, 'indexes')
_TUSAB_FORMAT_VERSION = 1


@router.get("/export/base-compartilhavel/{projeto}")
def export_base_compartilhavel(projeto: str):
    """Empacota textos processados + índice BM25 de um projeto num arquivo .tusab."""
    projeto = projeto.strip()
    if not projeto:
        return JSONResponse({"error": True, "message": "Nome do projeto não informado."})

    neural_path = os.path.join(NEURAL_DIR, projeto)
    index_path  = os.path.join(_INDEX_DIR, f"{projeto}.pkl")

    if not os.path.exists(neural_path):
        return JSONResponse({"error": True, "message": f"Projeto '{projeto}' não encontrado."})

    buf = io.BytesIO()
    chunk_count = 0

    with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zf:
        # Textos processados — apenas subpastas de conteúdo (sem arquivos originais grandes)
        for subdir in ('youtube', 'documents', 'texts', 'management'):
            subpath = os.path.join(neural_path, subdir)
            if not os.path.exists(subpath):
                continue
            for root, _, files in os.walk(subpath):
                for fname in files:
                    fpath = os.path.join(root, fname)
                    arcname = os.path.join('neural', projeto, os.path.relpath(fpath, neural_path))
                    zf.write(fpath, arcname)
                    if fname.endswith('.txt'):
                        chunk_count += 1

        # Índice BM25 serializado (aluno não precisa reindexar)
        if os.path.exists(index_path):
            zf.write(index_path, os.path.join('indexes', f"{projeto}.pkl"))

        # Manifest — somente_leitura: True marca a base como recebida/protegida no destino
        manifest = {
            "format": "tusab-base",
            "version": _TUSAB_FORMAT_VERSION,
            "projeto": projeto,
            "exportado_em": datetime.now().isoformat(),
            "chunks": chunk_count,
            "tem_indice": os.path.exists(index_path),
            "somente_leitura": True,
        }
        zf.writestr("manifest.json", json.dumps(manifest, ensure_ascii=False, indent=2))

    buf.seek(0)
    safe_name = projeto.replace(' ', '_')
    filename = f"{safe_name}.tusab"

    return StreamingResponse(
        buf,
        media_type="application/zip",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )


# ── Import de base compartilhável (.tusab) ────────────────────────────────────

@router.post("/import/base-compartilhavel")
async def import_base_compartilhavel(arquivo: UploadFile = File(...)):
    """Importa um arquivo .tusab — descompacta neural/ e indexes/ na pasta de dados."""
    if not arquivo.filename.endswith('.tusab'):
        return JSONResponse({"error": True, "message": "Arquivo deve ter extensão .tusab"})

    conteudo = await arquivo.read()
    buf = io.BytesIO(conteudo)

    try:
        with zipfile.ZipFile(buf, 'r') as zf:
            nomes = zf.namelist()

            # Valida manifest
            if 'manifest.json' not in nomes:
                return JSONResponse({"error": True, "message": "Arquivo .tusab inválido: manifest.json ausente."})

            manifest = json.loads(zf.read('manifest.json'))
            if manifest.get('format') != 'tusab-base':
                return JSONResponse({"error": True, "message": "Formato de arquivo não reconhecido."})

            projeto = manifest.get('projeto', '').strip()
            if not projeto:
                return JSONResponse({"error": True, "message": "Manifest sem nome de projeto."})

            # Extrai apenas neural/ e indexes/ — ignora qualquer outro caminho (segurança)
            for name in nomes:
                if not (name.startswith('neural/') or name.startswith('indexes/')):
                    continue
                # Proteção contra path traversal
                dest = os.path.join(DATA_DIR, name)
                dest = os.path.realpath(dest)
                if not dest.startswith(os.path.realpath(DATA_DIR)):
                    continue
                os.makedirs(os.path.dirname(dest), exist_ok=True)
                with zf.open(name) as src, open(dest, 'wb') as dst:
                    shutil.copyfileobj(src, dst)

            # Salva flag de readonly se manifest indica somente_leitura
            somente_leitura = manifest.get("somente_leitura", False)
            if somente_leitura:
                readonly_path = os.path.join(NEURAL_DIR, projeto, '_readonly.json')
                salvar_json_atomico(
                    {"somente_leitura": True, "importado_em": datetime.now().isoformat()},
                    readonly_path
                )

    except zipfile.BadZipFile:
        return JSONResponse({"error": True, "message": "Arquivo corrompido ou não é um .tusab válido."})

    return {
        "ok": True,
        "projeto": projeto,
        "chunks": manifest.get("chunks", 0),
        "tem_indice": manifest.get("tem_indice", False),
        "somente_leitura": manifest.get("somente_leitura", False),
        "message": f"Base '{projeto}' importada com sucesso.",
    }


# ── Status de readonly por projeto ───────────────────────────────────────────

@router.get("/export/readonly-status")
def readonly_status():
    """Retorna dict { projeto: bool } com quais projetos são somente leitura."""
    resultado = {}
    if not os.path.exists(NEURAL_DIR):
        return resultado
    for entry in os.scandir(NEURAL_DIR):
        if not entry.is_dir():
            continue
        readonly_path = os.path.join(entry.path, '_readonly.json')
        resultado[entry.name] = os.path.exists(readonly_path)
    return resultado
